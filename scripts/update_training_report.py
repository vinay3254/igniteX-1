from __future__ import annotations

import re
import shutil
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


DOC_PATH = Path(r"C:\Users\Admin\Desktop\igniteX\IgniteX_Complete_Training_Report.docx")


def insert_paragraph_before(paragraph, text: str, style: str | None = None):
    new_paragraph = paragraph.insert_paragraph_before(text)
    if style:
        new_paragraph.style = style
    return new_paragraph


def delete_paragraph(paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def main() -> None:
    backup_path = DOC_PATH.with_name(
        f"{DOC_PATH.stem}_backup_{datetime.now().strftime('%Y%m%d_%H%M%S')}{DOC_PATH.suffix}"
    )
    shutil.copy2(DOC_PATH, backup_path)

    doc = Document(DOC_PATH)
    paragraphs = doc.paragraphs

    title = paragraphs[0]
    title.style = doc.styles["Title"]
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = paragraphs[1]
    subtitle.style = doc.styles["Subtitle"]
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    program_details = paragraphs[2]
    program_details.text = (
        "Program Details: Trainee: Vinay GK | Training Reference Date: 14 March 2026 | "
        "Duration: 15 Training Days"
    )
    program_details.alignment = WD_ALIGN_PARAGRAPH.CENTER

    coverage = paragraphs[3]
    coverage.text = (
        "Coverage Areas: HTML5 | CSS3 | JavaScript ES6+ | React | Node.js | MongoDB | "
        "Full-Stack Development"
    )
    coverage.alignment = WD_ALIGN_PARAGRAPH.CENTER

    review_stamp = paragraphs[4]
    review_stamp.text = "Review Edition: Heading structure and review notes updated on 22 March 2026."
    review_stamp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    review_heading = paragraphs[9]
    review_heading.text = "How to Review This Report"
    review_heading.style = doc.styles["Heading 2"]

    review_intro = paragraphs[10]
    review_intro.text = (
        "Use Word's Navigation Pane to move through the report by heading level. "
        "The outline below defines how each section is organised for revision."
    )

    review_bullets = [
        "Heading 1 marks each training day or major closing section.",
        "Heading 2 marks a concept explanation, tool, project area, or key topic inside that section.",
        "Heading 3 marks detailed breakdowns such as line-by-line explanations or applied concept notes.",
        "Key Takeaways sections are short recall lists for quick revision before practice or interviews.",
        "The summary tables near the end work as a fast checklist across HTML, CSS, JavaScript, React, and backend topics.",
    ]

    first_day_index = next(
        i for i, para in enumerate(doc.paragraphs) if para.text.strip().startswith("DAY 1")
    )
    for para in list(doc.paragraphs[11:first_day_index]):
        delete_paragraph(para)

    anchor = next(para for para in doc.paragraphs if para.text.strip().startswith("DAY 1"))
    for bullet in review_bullets:
        insert_paragraph_before(anchor, bullet, "List Bullet")

    converted_day_titles = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if para.style.name == "Normal" and re.match(r"^DAY(S)?\b", text):
            para.style = doc.styles["Heading 1"]
            para.paragraph_format.page_break_before = True
            para.paragraph_format.keep_with_next = True
            converted_day_titles.append(text)

    for para in doc.paragraphs:
        if para.style.name.startswith("Heading"):
            para.paragraph_format.keep_with_next = True

    for para in doc.paragraphs:
        text = para.text.strip()
        if text in {
            "Complete Skills Reference — Summary Tables",
            "Conclusion",
        }:
            para.paragraph_format.page_break_before = True

    closing_line = doc.paragraphs[-1]
    closing_line.text = (
        "IgniteX Frontend Mastery Program | Complete Training Report | "
        "Review structure updated on 22 March 2026"
    )
    closing_line.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(DOC_PATH)

    print(f"Updated: {DOC_PATH}")
    print(f"Backup:  {backup_path}")
    print("Promoted day sections:", len(converted_day_titles))
    for title_text in converted_day_titles:
        print(f"- {title_text}")


if __name__ == "__main__":
    main()
